from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from tools.document_reader.models import Document, ReaderLimits

W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def parse_docx(path: Path, limits: ReaderLimits) -> Document:
    try:
        import docx  # type: ignore
    except ImportError:
        return _parse_docx_zip(path, limits)

    warnings: list[str] = []
    try:
        doc = docx.Document(path)
    except Exception as exc:
        return Document(path.name, "docx", warnings=[f"Could not parse DOCX: {exc}"])
    sections = []
    chars = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        kind = "heading" if style.lower().startswith("heading") else "paragraph"
        chars += len(text)
        if chars > limits.max_chars:
            warnings.append(f"DOCX text truncated to {limits.max_chars} characters.")
            break
        sections.append({"type": kind, "style": style, "text": text})
    tables = []
    doc_tables = doc.tables
    if len(doc_tables) > limits.max_tables:
        warnings.append(f"DOCX tables truncated to {limits.max_tables} from {len(doc_tables)}.")
        doc_tables = doc_tables[: limits.max_tables]
    for table_index, table in enumerate(doc_tables):
        rows = [[cell.text for cell in row.cells] for row in table.rows[: limits.table_preview_rows]]
        tables.append({"index": table_index, "headers": rows[0] if rows else [], "rows": rows[1:]})
    if not sections and not tables:
        warnings.append("DOCX appears empty.")
    return Document(path.name, "docx", metadata={}, sections=sections, tables=tables, warnings=warnings)


def _parse_docx_zip(path: Path, limits: ReaderLimits) -> Document:
    warnings = ["python-docx is not installed; using limited DOCX XML parser."]
    try:
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml")
    except zipfile.BadZipFile:
        return Document(path.name, "docx", warnings=["Could not parse DOCX: invalid or corrupt archive."])
    except Exception as exc:
        return Document(path.name, "docx", warnings=[f"Could not parse DOCX: {exc}"])
    root = ET.fromstring(xml)
    parent_map = {child: parent for parent in root.iter() for child in parent}
    tbl_tag = f"{{{W_NS['w']}}}tbl"

    def _has_ancestor_tag(elem: ET.Element, tag: str) -> bool:
        current = elem
        while current in parent_map:
            current = parent_map[current]
            if current.tag == tag:
                return True
        return False

    sections = []
    chars = 0
    # Only top-level paragraphs: paragraphs inside a table cell are already
    # represented in that table's row/cell text below, so including them here
    # too (via a recursive ".//w:p" search) would duplicate their content.
    for paragraph in root.findall(".//w:p", W_NS):
        if _has_ancestor_tag(paragraph, tbl_tag):
            continue
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", W_NS)).strip()
        if not text:
            continue
        style_node = paragraph.find(".//w:pStyle", W_NS)
        style = style_node.attrib.get(f"{{{W_NS['w']}}}val", "") if style_node is not None else ""
        kind = "heading" if style.lower().startswith("heading") else "paragraph"
        chars += len(text)
        if chars > limits.max_chars:
            warnings.append(f"DOCX text truncated to {limits.max_chars} characters.")
            break
        sections.append({"type": kind, "style": style, "text": text})
    tables = []
    # Only top-level tables: a table nested inside another table's cell would
    # otherwise be extracted twice (once via the parent cell's flattened text,
    # once as its own entry via a recursive ".//w:tbl" search).
    top_level_tables = [
        table for table in root.findall(".//w:tbl", W_NS) if not _has_ancestor_tag(table, tbl_tag)
    ]
    if len(top_level_tables) > limits.max_tables:
        warnings.append(f"DOCX tables truncated to {limits.max_tables} from {len(top_level_tables)}.")
        top_level_tables = top_level_tables[: limits.max_tables]
    for table_index, table in enumerate(top_level_tables):
        rows = []
        # Direct-child "w:tr" only: a nested table inside a cell also has its
        # own "w:tr" descendants, and ".//w:tr" would pull those rows into
        # this table too.
        for row in table.findall("w:tr", W_NS)[: limits.table_preview_rows]:
            rows.append(["".join(t.text or "" for t in cell.findall(".//w:t", W_NS)) for cell in row.findall("w:tc", W_NS)])
        tables.append({"index": table_index, "headers": rows[0] if rows else [], "rows": rows[1:]})
    if not sections and not tables:
        warnings.append("DOCX appears empty.")
    return Document(path.name, "docx", sections=sections, tables=tables, warnings=warnings)

